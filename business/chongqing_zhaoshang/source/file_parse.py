from docx import Document
import os
#from my_log import logger

def get_files_in_directory(directory):
    """
    获取指定目录下的所有文件（不包含子目录中的文件）
    :param directory: 相对或绝对路径
    :return: 文件列表
    """
    # 确保路径是绝对路径
    directory = os.path.abspath(directory)

    # 检查路径是否存在
    if not os.path.exists(directory):
        raise ValueError(f"The directory {directory} does not exist.")

    # 列出目录中的所有条目
    entries = os.listdir(directory)

    # 过滤出文件
    files = [entry for entry in entries if os.path.isfile(os.path.join(directory, entry))]

    return files
def file_parse():

    file_content_list = {}
    directory = '../files'  # 相对路径
    # 确保路径是绝对路径
    directory = os.path.abspath(directory)

    # 检查路径是否存在
    if not os.path.exists(directory):
        raise ValueError(f"The directory {directory} does not exist.")

    # 列出目录中的所有条目
    entries = os.listdir(directory)
    for entry in entries:
        file = os.path.join(directory, entry)

        print(file)
        doc = Document(file)
        # 定义一个空列表，用于存储切分后的内容
        content_list = []

        # 定义一个空字符串，用于存储当前章节的内容
        current_content = ""

        # 遍历文档中的每个段落
        for paragraph in doc.paragraphs:
            # 如果当前段落是标题，则将当前章节的内容添加到内容列表中，并重新开始一个新章节
            if paragraph.style.name.startswith("Heading"):
                if current_content != "":
                    content_list.append(current_content)
                    current_content = ""
            # 将当前段落的文本添加到当前章节的内容中
            current_content += f"{paragraph.text}\n"

        # 将最后一个章节的内容添加到内容列表中
        if current_content != "":
            content_list.append(current_content)

        # 输出切分后的内容
        for i, content in enumerate(content_list):
            #print(f"Chapter {i+1}:\n{content}\n")
            cont = file_content_list.get(i, None)
            if cont:
                cont += content
                file_content_list[i] = cont
            else:
                file_content_list[i] = content

    for key, file_cont in file_content_list.items():
        print(f"file_pars= ekey:{key} file_cont :\n{file_cont}\n")
        #logger.info(f"file_pars= ekey:{key} file_cont :\n{file_cont}\n")
    return file_content_list



# # 定义一个空列表，用于存储分割后的章节
# chapters = []
#
# # 定义一个空字符串，用于存储当前章节的标题
# current_title = ""
#
# # 定义一个空字符串，用于存储当前章节的内容
# current_content = ""
#
# # 遍历文档中的每个段落
# for paragraph in doc.paragraphs:
#     # 获取当前段落的样式
#     style = paragraph.style.name
#     print(f"style={style} paragraph.text={paragraph.text}")
#
#     # 如果当前段落是标题，则将当前章节的标题和内容添加到章节列表中，并重新开始一个新章节
#     if style.startswith("Heading"):
#         # 获取当前标题的级别
#         level = int(style.replace("Heading", ""))
#         # 如果当前标题是父级标题，则将当前章节的标题和内容添加到章节列表中，并重新开始一个新章节
#         if level == 1:
#             if current_title != "":
#                 chapters.append({"title": current_title, "content": current_content})
#             current_title = paragraph.text
#             current_content = ""
#         # 如果当前标题是子级标题，则将当前段落的文本添加到当前章节的内容中
#         elif level > 1:
#             current_content += f"{paragraph.text}\n"
#     # 如果当前段落不是标题，则将当前段落的文本添加到当前章节的内容中
#     else:
#         current_content += f"{paragraph.text}\n"
#
# # 将最后一个章节的标题和内容添加到章节列表中
# chapters.append({"title": current_title, "content": current_content})
#
# # 输出分割后的章节
# for i, chapter in enumerate(chapters):
#     print(f"Chapter {i+1} - {chapter['title']}:\n{chapter['content']}\n")


if __name__ == '__main__':
    file_parse()